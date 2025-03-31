"""
このファイルは、最初の画面読み込み時にのみ実行される初期化処理が記述されたファイルです。
"""

############################################################
# ライブラリの読み込み
############################################################
import os
import logging
from logging.handlers import TimedRotatingFileHandler
from uuid import uuid4
import sys
import unicodedata
from dotenv import load_dotenv
import streamlit as st
from docx import Document
from langchain_community.document_loaders import WebBaseLoader
from langchain.text_splitter import CharacterTextSplitter
from langchain_openai import OpenAIEmbeddings
from langchain_community.vectorstores import Chroma
import constants as ct


############################################################
# 設定関連
############################################################
# 「.env」ファイルで定義した環境変数の読み込み
load_dotenv()


############################################################
# 関数定義
############################################################

def initialize():
    """
    画面読み込み時に実行する初期化処理
    """
    # 初期化データの用意
    initialize_session_state()
    # ログ出力用にセッションIDを生成
    initialize_session_id()
    # ログ出力の設定
    initialize_logger()
    # RAGのRetrieverを作成
    initialize_retriever()


def initialize_logger():
    """
    ログ出力の設定
    """
    # 指定のログフォルダが存在すれば読み込み、存在しなければ新規作成
    os.makedirs(ct.LOG_DIR_PATH, exist_ok=True)
    
    # 引数に指定した名前のロガー（ログを記録するオブジェクト）を取得
    # 再度別の箇所で呼び出した場合、すでに同じ名前のロガーが存在していれば読み込む
    logger = logging.getLogger(ct.LOGGER_NAME)

    # すでにロガーにハンドラー（ログの出力先を制御するもの）が設定されている場合、同じログ出力が複数回行われないよう処理を中断する
    if logger.hasHandlers():
        return

    # 1日単位でログファイルの中身をリセットし、切り替える設定
    log_handler = TimedRotatingFileHandler(
        os.path.join(ct.LOG_DIR_PATH, ct.LOG_FILE),
        when="D",
        encoding="utf8"
    )
    # 出力するログメッセージのフォーマット定義
    # - 「levelname」: ログの重要度（INFO, WARNING, ERRORなど）
    # - 「asctime」: ログのタイムスタンプ（いつ記録されたか）
    # - 「lineno」: ログが出力されたファイルの行番号
    # - 「funcName」: ログが出力された関数名
    # - 「session_id」: セッションID（誰のアプリ操作か分かるように）
    # - 「message」: ログメッセージ
    formatter = logging.Formatter(
        f"[%(levelname)s] %(asctime)s line %(lineno)s, in %(funcName)s, session_id={st.session_state.session_id}: %(message)s"
    )

    # 定義したフォーマッターの適用
    log_handler.setFormatter(formatter)

    # ログレベルを「INFO」に設定
    logger.setLevel(logging.INFO)

    # 作成したハンドラー（ログ出力先を制御するオブジェクト）を、
    # ロガー（ログメッセージを実際に生成するオブジェクト）に追加してログ出力の最終設定
    logger.addHandler(log_handler)


def initialize_session_id():
    """
    セッションIDの作成
    """
    if "session_id" not in st.session_state:
        # ランダムな文字列（セッションID）を、ログ出力用に作成
        st.session_state.session_id = uuid4().hex


def initialize_retriever():
    """
    画面読み込み時にRAGのRetriever（ベクターストアから検索するオブジェクト）を作成
    """
    # ロガーを読み込むことで、後続の処理中に発生したエラーなどがログファイルに記録される
    logger = logging.getLogger(ct.LOGGER_NAME)

    # すでにRetrieverが作成済みの場合、後続の処理を中断
    if "retriever" in st.session_state:
        return
    
    # RAGの参照先となるデータソースの読み込み
    docs_all = load_data_sources()
    
    # 社員名簿.csvファイルを特別に処理する
    employee_docs = process_employee_csv()
    if employee_docs:
        # 既存のドキュメントリストから社員名簿.csvに関するドキュメントを削除
        docs_all = [doc for doc in docs_all if "社員名簿.csv" not in doc.metadata.get("source", "")]
        # 新しく作成した社員名簿ドキュメントを追加
        docs_all.extend(employee_docs)

    # OSがWindowsの場合、Unicode正規化と、cp932（Windows用の文字コード）で表現できない文字を除去
    for doc in docs_all:
        doc.page_content = adjust_string(doc.page_content)
        for key in doc.metadata:
            doc.metadata[key] = adjust_string(doc.metadata[key])
    
    # 埋め込みモデルの用意
    embeddings = OpenAIEmbeddings()
    
    # 通常のドキュメント用のチャンク分割設定
    text_splitter = CharacterTextSplitter(
        chunk_size=ct.RAG_CHUNK_SIZE,
        chunk_overlap=ct.RAG_CHUNK_OVERLAP,
        separator="\n"
    )

    # 社員名簿以外のドキュメントをチャンク分割
    normal_docs = [doc for doc in docs_all if not doc.metadata.get("is_employee_list", False)]
    employee_docs = [doc for doc in docs_all if doc.metadata.get("is_employee_list", False)]
    
    # 通常のドキュメントのみチャンク分割
    splitted_normal_docs = text_splitter.split_documents(normal_docs) if normal_docs else []
    
    # 全てのドキュメントを結合（社員名簿はチャンク分割しない）
    all_docs = splitted_normal_docs + employee_docs

    # ベクターストアの作成
    db = Chroma.from_documents(all_docs, embedding=embeddings)

    # ベクターストアを検索するRetrieverの作成
    # 変更日時: 2025-03-30
    # 変更内容: 関連ドキュメントの数を3から5に変更し、マジックナンバーを定数化
    st.session_state.retriever = db.as_retriever(search_kwargs={"k": ct.RAG_RETRIEVER_TOP_K})


def process_employee_csv():
    """
    社員名簿.csvファイルを処理し、特別なドキュメントを作成する
    
    Returns:
        社員名簿から作成したドキュメントのリスト
    """
    import pandas as pd
    from langchain.schema import Document
    
    csv_path = "data/社員について/社員名簿.csv"
    if not os.path.exists(csv_path):
        return []
    
    try:
        # CSVファイルを読み込む
        df = pd.read_csv(csv_path)
        
        # 人事部の社員をフィルタリング
        hr_employees = df[df['部署'] == '人事部']
        
        # 人事部の社員一覧を作成（検索しやすいフォーマット）
        hr_content = """
# 人事部の社員一覧

以下は人事部に所属する全ての社員の情報です。

"""
        for _, row in hr_employees.iterrows():
            hr_content += f"""
## {row['氏名（フルネーム）']} ({row['社員ID']})
- **役職**: {row['役職']}
- **性別**: {row['性別']}
- **年齢**: {row['年齢']}
- **メールアドレス**: {row['メールアドレス']}
- **従業員区分**: {row['従業員区分']}
- **入社日**: {row['入社日']}
- **スキルセット**: {row['スキルセット']}
- **保有資格**: {row['保有資格']}
- **学歴**: {row['大学名']} {row['学部・学科']} ({row['卒業年月日']}卒業)

"""
        
        # 各部署ごとの社員一覧を作成
        department_docs = []
        departments = df['部署'].unique()
        
        for dept in departments:
            dept_employees = df[df['部署'] == dept]
            dept_content = f"""
# {dept}の社員一覧

以下は{dept}に所属する全ての社員の情報です。

"""
            for _, row in dept_employees.iterrows():
                dept_content += f"""
## {row['氏名（フルネーム）']} ({row['社員ID']})
- **役職**: {row['役職']}
- **性別**: {row['性別']}
- **年齢**: {row['年齢']}
- **メールアドレス**: {row['メールアドレス']}
- **従業員区分**: {row['従業員区分']}
- **入社日**: {row['入社日']}
- **スキルセット**: {row['スキルセット']}
- **保有資格**: {row['保有資格']}
- **学歴**: {row['大学名']} {row['学部・学科']} ({row['卒業年月日']}卒業)

"""
            
            # 部署ごとのドキュメントを作成
            dept_doc = Document(
                page_content=dept_content,
                metadata={
                    "source": csv_path,
                    "department": dept,
                    "is_employee_list": True,
                    "content_type": "department_list"
                }
            )
            department_docs.append(dept_doc)
        
        # 人事部の社員一覧ドキュメントを作成
        hr_doc = Document(
            page_content=hr_content,
            metadata={
                "source": csv_path,
                "department": "人事部",
                "is_employee_list": True,
                "content_type": "hr_department_list",
                "description": "人事部の社員一覧情報"
            }
        )
        
        # 特別なクエリ用のドキュメントを作成
        query_doc = Document(
            page_content="""
人事部の社員一覧:
このドキュメントには人事部に所属する全ての社員の詳細情報が含まれています。
「人事部に所属している従業員情報を一覧化して」という質問に対応するためのドキュメントです。
""",
            metadata={
                "source": csv_path,
                "is_employee_list": True,
                "content_type": "query_helper",
                "keywords": "人事部,社員一覧,従業員情報"
            }
        )
        
        # 全てのドキュメントを結合
        all_docs = [hr_doc, query_doc] + department_docs
        return all_docs
        
    except Exception as e:
        print(f"社員名簿の処理中にエラーが発生しました: {e}")
        return []


def initialize_session_state():
    """
    初期化データの用意
    """
    if "messages" not in st.session_state:
        # 「表示用」の会話ログを順次格納するリストを用意
        st.session_state.messages = []
        # 「LLMとのやりとり用」の会話ログを順次格納するリストを用意
        st.session_state.chat_history = []


def load_data_sources():
    """
    RAGの参照先となるデータソースの読み込み

    Returns:
        読み込んだ通常データソース
    """
    # データソースを格納する用のリスト
    docs_all = []
    # ファイル読み込みの実行（渡した各リストにデータが格納される）
    recursive_file_check(ct.RAG_TOP_FOLDER_PATH, docs_all)

    web_docs_all = []
    # ファイルとは別に、指定のWebページ内のデータも読み込み
    # 読み込み対象のWebページ一覧に対して処理
    for web_url in ct.WEB_URL_LOAD_TARGETS:
        # 指定のWebページを読み込み
        loader = WebBaseLoader(web_url)
        web_docs = loader.load()
        # for文の外のリストに読み込んだデータソースを追加
        web_docs_all.extend(web_docs)
    # 通常読み込みのデータソースにWebページのデータを追加
    docs_all.extend(web_docs_all)

    return docs_all


def recursive_file_check(path, docs_all):
    """
    RAGの参照先となるデータソースの読み込み

    Args:
        path: 読み込み対象のファイル/フォルダのパス
        docs_all: データソースを格納する用のリスト
    """
    # パスがフォルダかどうかを確認
    if os.path.isdir(path):
        # フォルダの場合、フォルダ内のファイル/フォルダ名の一覧を取得
        files = os.listdir(path)
        # 各ファイル/フォルダに対して処理
        for file in files:
            # ファイル/フォルダ名だけでなく、フルパスを取得
            full_path = os.path.join(path, file)
            # フルパスを渡し、再帰的にファイル読み込みの関数を実行
            recursive_file_check(full_path, docs_all)
    else:
        # パスがファイルの場合、ファイル読み込み
        file_load(path, docs_all)


def file_load(path, docs_all):
    """
    ファイル内のデータ読み込み

    Args:
        path: ファイルパス
        docs_all: データソースを格納する用のリスト
    """
    # ファイルの拡張子を取得
    file_extension = os.path.splitext(path)[1]
    # ファイル名（拡張子を含む）を取得
    file_name = os.path.basename(path)

    # 想定していたファイル形式の場合のみ読み込む
    if file_extension in ct.SUPPORTED_EXTENSIONS:
        # 社員名簿.csvファイルの場合はprocess_employee_csv関数で処理するためスキップ
        if path.endswith("社員について/社員名簿.csv"):
            # 何もしない（process_employee_csv関数で処理される）
            pass
        else:
            # 通常のファイル読み込み処理
            loader = ct.SUPPORTED_EXTENSIONS[file_extension](path)
            docs = loader.load()
            docs_all.extend(docs)


def adjust_string(s):
    """
    Windows環境でRAGが正常動作するよう調整
    
    Args:
        s: 調整を行う文字列
    
    Returns:
        調整を行った文字列
    """
    # 調整対象は文字列のみ
    if type(s) is not str:
        return s

    # OSがWindowsの場合、Unicode正規化と、cp932（Windows用の文字コード）で表現できない文字を除去
    if sys.platform.startswith("win"):
        s = unicodedata.normalize('NFC', s)
        s = s.encode("cp932", "ignore").decode("cp932")
        return s
    
    # OSがWindows以外の場合はそのまま返す
    return s